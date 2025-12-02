"""
Generate a concise Turkish Word document describing the MongoDB homework.
This script creates `MongoDB_Homework_[STUDENTNUMBER]_TR.docx` in the project root.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime

student_number = input("Öğrenci numaranızı girin (örnek: STU001): ").strip() or "STU000"
student_name = input("Adınızı girin: ").strip() or "İsim Soyisim"

doc = Document()

# Title
h = doc.add_heading('MongoDB & Python Entegrasyonu - Kısa Rapor', level=0)
h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
h.runs[0].font.size = Pt(16)

# Meta
meta = doc.add_paragraph()
meta.add_run('Öğrenci Numarası: ').bold = True
meta.add_run(student_number + "\n")
meta.add_run('Ad: ').bold = True
meta.add_run(student_name + "\n")
meta.add_run('Tarih: ').bold = True
meta.add_run(datetime.datetime.now().strftime('%d/%m/%Y') + "\n")

# Professor instruction reminder
p = doc.add_paragraph()
p.add_run('Not: Bölüm konuları öğrenci numaralarına göre dağıtılmıştır; herkes kendi verilen konu üzerinde çalışacaktır. ')
p.add_run('Çalışmada farklı konu seçilmesi yasaktır.').italic = True

# Purpose
doc.add_heading('Amaç', level=1)
doc.add_paragraph('Bu çalışma PyMongo kullanarak MongoDB ile Python entegrasyonunu göstermeyi amaçlar. ' 
                  'Kısaca CRUD işlemleri, aggregate(), find_one(), count_documents() ve verinin Pandas DataFrame\'e aktarımı gösterilmiştir.')

# Short Examples heading
doc.add_heading('Kısa Örnekler ve Sorgu Blokları', level=1)

# CREATE
doc.add_heading('1) CREATE (Ekleme)', level=2)
code = (
    "# Tek doküman ekleme\n"
    "students.insert_one({\n"
    "    'name': 'Ali', 'student_id': 'STU010', 'dept': 'CS', 'gpa': 3.8\n"
    "})\n\n"
    "# Çoklu ekleme\n"
    "students.insert_many([{'name':'Ayşe','student_id':'STU011','dept':'ENG'},{'name':'Mehmet','student_id':'STU012','dept':'MATH'}])"
)
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# READ
doc.add_heading('2) READ (find, find_one)', level=2)
code = (
    "# Tek eşleşen doküman\n"
    "doc = students.find_one({'dept':'CS'})\n"
    "# Tüm dokümanlar\n"
    "for s in students.find():\n    print(s)\n\n"
    "# Projection (alan seçimi)\n"
    "students.find({}, {'name':1,'student_id':1,'_id':0})"
)
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# COUNT
doc.add_heading('3) Kayıt Sayımı (count_documents)', level=2)
code = (
    "# Tüm kayıtlar\n"
    "total = students.count_documents({})\n"
    "# Bölüme göre sayma\n"
    "cs_count = students.count_documents({'dept':'CS'})"
)
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# AGGREGATE
doc.add_heading('4) Aggregation (aggregate())', level=2)
code = (
    "# Bölüme göre ortalama GPA\n"
    "pipeline = [\n"
    "  {'$group': {'_id':'$dept','avg_gpa':{'$avg':'$gpa'},'count':{'$sum':1}}},\n"
    "  {'$sort': {'avg_gpa':-1}}\n"
    "]\n"
    "for r in students.aggregate(pipeline): print(r)"
)
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# UPDATE
doc.add_heading('5) UPDATE (Güncelleme)', level=2)
code = (
    "# Tek doküman güncelle\n"
    "students.update_one({'student_id':'STU010'},{'$set':{'gpa':3.9}})\n\n"
    "# Çoklu güncelle (ör: CS öğrencilerinin GPA artışı)\n"
    "students.update_many({'dept':'CS'},{'$inc':{'gpa':0.1}})"
)
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# DELETE
doc.add_heading('6) DELETE (Silme)', level=2)
code = (
    "# Tek silme\n"
    "students.delete_one({'student_id':'STU999'})\n\n"
    "# Koşula göre birden çok silme\n"
    "students.delete_many({'gpa':{'$lt':3.0}})"
)
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# DATAFRAME
doc.add_heading('7) DataFrame (Pandas)', level=2)
code = (
    "import pandas as pd\n"
    "data = list(students.find({}, {'_id':0}))\n"
    "df = pd.DataFrame(data)\n"
    "print(df.head())\n"
    "df.to_csv('students.csv', index=False)"
)
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# Brief notes
doc.add_heading('Kısa Notlar', level=1)
doc.add_paragraph('- MongoDB bağlantısı: mongodb://localhost:27017\n- `crud_examples.py` gerçek MongoDB gerektirir; `crud_demo.py` sunucusuz test içindir.\n- Hazırlanan belge öğrenci numarası ile teslim edilecektir.')

# Save
output = f"/home/abod/Workspace/Projects/Python-MongoDB-Integration/MongoDB_Homework_{student_number}_TR.docx"
doc.save(output)
print(f"Oluşturuldu: {output}")
