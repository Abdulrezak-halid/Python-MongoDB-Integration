"""
Generate Word Document for MongoDB Homework Submission
This script creates a professional Word document containing:
- Assignment purpose
- Query blocks with explanations
- MongoDB operations documentation
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime

# Prompt user for student information
student_number = input("Enter your student number: ").strip()
student_name = input("Enter your full name: ").strip()

# Create Document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading("MongoDB CRUD Operations with PyMongo", level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_format = title.runs[0].font
title_format.size = Pt(16)
title_format.bold = True

# Course Info
course_info = doc.add_paragraph()
course_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = course_info.add_run(f"Subject: Python & MongoDB Integration\nDistributed Database Systems\n")
run.font.size = Pt(11)
run.font.italic = True

# Student Info
student_info = doc.add_paragraph()
student_info.add_run(f"Student Number: ").bold = True
student_info.add_run(student_number)
student_info.add_run(f"\nStudent Name: ").bold = True
student_info.add_run(student_name)
student_info.add_run(f"\nDate: ").bold = True
student_info.add_run(datetime.datetime.now().strftime("%d/%m/%Y"))

doc.add_paragraph()

# Assignment Purpose
doc.add_heading("1. Assignment Purpose", level=1)
purpose = doc.add_paragraph(
    "The objective of this assignment is to demonstrate proficiency in MongoDB database operations "
    "using PyMongo Python driver. This includes implementing CRUD (Create, Read, Update, Delete) "
    "operations, utilizing aggregation pipelines for complex data analysis, counting documents with "
    "filters, and converting MongoDB data into Pandas DataFrames for further analysis and export."
)
purpose_points = [
    "Master PyMongo driver for MongoDB connectivity",
    "Implement all CRUD operations (Create, Read, Update, Delete)",
    "Use MongoDB aggregation pipeline for data analysis",
    "Apply filtering and projection in queries",
    "Work with count_documents() for data statistics",
    "Convert MongoDB documents to Pandas DataFrames",
    "Export data to CSV and Excel formats"
]
for point in purpose_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph()

# Query Blocks Section
doc.add_heading("2. MongoDB Query Blocks", level=1)

# Query Block 1: INSERT
doc.add_heading("Query Block 1: CREATE (INSERT) Operations", level=2)
doc.add_paragraph("Purpose: Add new student records to the MongoDB collection").runs[0].italic = True

code1 = """# Insert single document
student = {
    "name": "Student Name",
    "student_id": "STU001",
    "age": 21,
    "dept": "CS",
    "gpa": 3.8
}
result = students.insert_one(student)
print(f"Inserted ID: {result.inserted_id}")

# Insert multiple documents
students_list = [
    {"name": "Student 1", "student_id": "STU001", "dept": "CS"},
    {"name": "Student 2", "student_id": "STU002", "dept": "ENG"}
]
result = students.insert_many(students_list)
print(f"Inserted {len(result.inserted_ids)} documents")"""

paragraph = doc.add_paragraph(code1)
for run in paragraph.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_paragraph("Explanation: insert_one() adds a single document, insert_many() adds multiple documents. "
                  "The method returns an InsertResult object containing the inserted document IDs.")

doc.add_paragraph()

# Query Block 2: FIND and FIND_ONE
doc.add_heading("Query Block 2: READ Operations (find, find_one)", level=2)
doc.add_paragraph("Purpose: Retrieve documents from the collection").runs[0].italic = True

code2 = """# Find first matching document
first_student = students.find_one({"dept": "CS"})
print(f"Found: {first_student['name']}")

# Find all documents with filter
cs_students = students.find({"dept": "CS"})
for student in cs_students:
    print(student['name'])

# Find with projection (select specific fields)
result = students.find(
    {"dept": "CS"},
    {"name": 1, "student_id": 1, "_id": 0}
)
for doc in result:
    print(doc)

# Find with complex filter
high_gpa = students.find({"gpa": {"$gt": 3.7}})"""

paragraph = doc.add_paragraph(code2)
for run in paragraph.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_paragraph("Explanation: find_one() returns first matching document or None. find() returns a cursor "
                  "for iteration. Projection (second parameter) selects fields (1=include, 0=exclude).")

doc.add_paragraph()

# Query Block 3: COUNT
doc.add_heading("Query Block 3: COUNT_DOCUMENTS() Operations", level=2)
doc.add_paragraph("Purpose: Count documents matching specific criteria").runs[0].italic = True

code3 = """# Count all documents
total = students.count_documents({})
print(f"Total students: {total}")

# Count with filter
cs_count = students.count_documents({"dept": "CS"})
print(f"CS students: {cs_count}")

# Count with complex condition
high_gpa_count = students.count_documents({"gpa": {"$gt": 3.7}})
print(f"Students with GPA > 3.7: {high_gpa_count}")

# Count by multiple conditions
count = students.count_documents({
    "dept": "CS",
    "gpa": {"$gte": 3.5}
})"""

paragraph = doc.add_paragraph(code3)
for run in paragraph.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_paragraph("Explanation: count_documents() returns the number of documents matching the filter. "
                  "Use empty {} to count all documents. Supports comparison operators like $gt, $gte, etc.")

doc.add_paragraph()

# Query Block 4: UPDATE
doc.add_heading("Query Block 4: UPDATE Operations", level=2)
doc.add_paragraph("Purpose: Modify existing documents").runs[0].italic = True

code4 = """# Update single document
result = students.update_one(
    {"student_id": "STU001"},
    {"$set": {"age": 22, "gpa": 3.9}}
)
print(f"Modified: {result.modified_count}")

# Update multiple documents
result = students.update_many(
    {"dept": "CS"},
    {"$inc": {"gpa": 0.1}}
)

# Upsert (update or insert)
students.update_one(
    {"student_id": "STU999"},
    {"$set": {"name": "New", "dept": "CS"}},
    upsert=True
)"""

paragraph = doc.add_paragraph(code4)
for run in paragraph.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_paragraph("Explanation: update_one() modifies first matching document, update_many() modifies all matching. "
                  "$set replaces fields, $inc increments values. upsert=True creates document if not found.")

doc.add_paragraph()

# Query Block 5: AGGREGATION
doc.add_heading("Query Block 5: AGGREGATION Pipeline", level=2)
doc.add_paragraph("Purpose: Perform complex data analysis and transformations").runs[0].italic = True

code5 = """# Group by department and calculate statistics
pipeline = [
    {
        "$group": {
            "_id": "$dept",
            "avg_gpa": {"$avg": "$gpa"},
            "count": {"$sum": 1},
            "max_gpa": {"$max": "$gpa"}
        }
    },
    {"$sort": {"avg_gpa": -1}}
]
for result in students.aggregate(pipeline):
    print(result)

# Match, project and sort
pipeline = [
    {"$match": {"gpa": {"$gte": 3.7}}},
    {"$project": {"name": 1, "gpa": 1, "_id": 0}},
    {"$sort": {"gpa": -1}},
    {"$limit": 5}
]"""

paragraph = doc.add_paragraph(code5)
for run in paragraph.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_paragraph("Explanation: Aggregation pipeline processes documents through stages. "
                  "$match filters, $group summarizes, $project selects fields, $sort orders results.")

doc.add_paragraph()

# Query Block 6: DELETE
doc.add_heading("Query Block 6: DELETE Operations", level=2)
doc.add_paragraph("Purpose: Remove documents from collection").runs[0].italic = True

code6 = """# Delete single document
result = students.delete_one({"student_id": "STU999"})
print(f"Deleted: {result.deleted_count}")

# Delete multiple documents
result = students.delete_many({"gpa": {"$lt": 3.5}})
print(f"Deleted {result.deleted_count} students with low GPA")

# Delete all documents (with empty filter)
# students.delete_many({})"""

paragraph = doc.add_paragraph(code6)
for run in paragraph.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_paragraph("Explanation: delete_one() removes first matching document, delete_many() removes all matching. "
                  "Returns DeleteResult with deleted_count property.")

doc.add_paragraph()

# Query Block 7: DataFrame
doc.add_heading("Query Block 7: DATAFRAME Conversion", level=2)
doc.add_paragraph("Purpose: Convert MongoDB data to Pandas for analysis and export").runs[0].italic = True

code7 = """# Convert MongoDB data to DataFrame
import pandas as pd

data = list(students.find({}, {"_id": 0}))
df = pd.DataFrame(data)

# Display DataFrame
print(df)
print(f"Shape: {df.shape}")

# Analysis
print(f"Average GPA: {df['gpa'].mean()}")
print(f"Students by department:")
print(df['dept'].value_counts())

# Export to CSV
df.to_csv('students.csv', index=False)

# Export to Excel
df.to_excel('students.xlsx', index=False, sheet_name='Students')"""

paragraph = doc.add_paragraph(code7)
for run in paragraph.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_paragraph("Explanation: list(students.find()) retrieves all documents. pd.DataFrame() converts to table format. "
                  "Enables statistical analysis and export to multiple formats (CSV, Excel, JSON, etc.).")

doc.add_paragraph()

# Connection Info
doc.add_heading("3. MongoDB Connection Setup", level=1)
connection_code = """from pymongo import MongoClient

# Connect to MongoDB
client = MongoClient("mongodb://localhost:27017")
db = client["school"]
students = db["students"]"""

paragraph = doc.add_paragraph(connection_code)
for run in paragraph.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_paragraph("Ensure MongoDB server is running on localhost:27017 before executing queries.")

doc.add_paragraph()

# Key Concepts
doc.add_heading("4. Key MongoDB Concepts", level=1)
concepts = {
    "Collection": "Similar to a database table, contains multiple documents",
    "Document": "JSON-like data structure, similar to a row in SQL",
    "Field": "Key-value pair within a document, similar to a column",
    "Query Operators": "$gt, $gte, $lt, $lte, $eq, $ne for comparisons",
    "Update Operators": "$set (replace), $inc (increment), $push (add to array)",
    "Aggregation Stages": "$match, $group, $project, $sort, $limit, $skip"
}

for concept, definition in concepts.items():
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(concept + ": ").bold = True
    p.add_run(definition)

doc.add_paragraph()

# Conclusion
doc.add_heading("5. Conclusion", level=1)
doc.add_paragraph(
    "This assignment demonstrates comprehensive MongoDB operations using PyMongo. The implemented "
    "queries cover all fundamental operations including CRUD, aggregation, and data export. "
    "Students can extend these examples for their specific use cases and database requirements."
)

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
footer.add_run(f"\n\nSubmitted: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")
footer.runs[0].italic = True
footer.runs[0].font.size = Pt(9)

# Save document
output_file = f"/home/abod/Workspace/Projects/Python-MongoDB-Integration/MongoDB_Homework_{student_number}.docx"
doc.save(output_file)
print(f"\n✓ Word document created successfully!")
print(f"📄 Saved to: {output_file}")
